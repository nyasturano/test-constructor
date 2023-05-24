import random
import re

from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import utils


class Variant:
    pass

    def __init__(self, t_id, src_path):
        self.src = src_path
        self.id = t_id
        self.test = Document()
        self.tasks = []
        self.key = Document()
        # make margins smaller for output test
        sections = self.test.sections
        for section in sections:
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        # enter header
        p = self.test.add_paragraph(f'Вариант {int(t_id) + 1}')
        p.runs[0].bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_task(self, task):
        # add task to list
        self.tasks.append(task)

        # get root element of test
        root = self.test._element.getchildren()[0]

        # add all pictures to the start of a test
        img_p = []
        for i in range(0, len(task.img)):
            p = self.test.add_paragraph()
            run = p.add_run()
            run.add_picture(f'{self.src}/img/{task.id}_{i}.png', width=Mm(65))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            img_p.append(p)

        # add task nodes to a test + replace matching images if they are
        for n in task.nodes:
            root.append(n)
            if utils.contains_picture(n):
                root.replace(n, img_p[0]._element)
                img_p.pop(0)

        # reorder task answers
        task.compile(self.test)
        # add answers to test
        root.append(task.ans)

    def form_key(self):
        # create table with 2 rows and as many columns as tasks
        table = self.key.add_table(2, len(self.tasks))
        table.style = 'Table Grid'

        # fill table with key values and add some style
        head_row = table.rows[0].cells
        for i, col in enumerate(head_row):
            run = col.paragraphs[0].add_run(f'{i + 1} ({int(self.tasks[i].id)})')
            run.bold = False
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
            col.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        key_row = table.rows[1].cells
        for i, task in enumerate(self.tasks):
            run = key_row[i].paragraphs[0].add_run(f'{task.right + 1}')
            run.bold = False
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
        return table

    def save_test(self, out_path):
        # numeration
        num = 1
        for p in self.test.paragraphs:
            if p.text is not None and bool(re.match(r'Задача\s\d+$', p.text)):
                p.text = f'№{num}'
                run = p.runs[0]
                run.bold = True
                run.font.name = 'Times New Roman'
                num += 1

        self.test.save(f'{out_path}/variant_{self.id + 1}.docx')

    def save_key(self, out_path):
        self.key.save(f'{out_path}/key_{self.id + 1}.docx')
