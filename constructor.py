from task import Task
from variant import Variant
import docx
import random
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Inches


def make_variants(v_count, th_mask, pr_mask, src_path='src', out_path='out'):

    # create document for keys
    key = docx.Document()
    sections = key.sections
    for section in sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    # root element of key document
    key_root = key._element.getchildren()[0]

    for v_c in range(0, int(v_count)):
        # open theory and practice source
        t_doc = docx.Document(f'{src_path}/t_source.docx')
        p_doc = docx.Document(f'{src_path}/p_source.docx')

        # create variant document
        variant = Variant(v_c, src_path)

        th_tasks = []
        pr_tasks = []

        # fetch tasks by headers
        for p in t_doc.paragraphs:
            if p.text is not None and bool(re.match(r'Задача\s\d+$', p.text)):
                th_tasks.append(Task(re.search(r'\d+', p.text).group(0)))
            th_tasks[-1].add_node(p._element)

        for p in p_doc.paragraphs:
            if p.text is not None and bool(re.match(r'Задача\s\d+$', p.text)):
                pr_tasks.append(Task(re.search(r'\d+', p.text).group(0), solution=True))
            pr_tasks[-1].add_node(p._element)

        # choose random theory task's indices
        indices = []
        for i in range(0, 5):
            indices.append(random.randrange(i * 10, (i + 1) * 10))

        # add targeted tasks to a variant document
        for idx, idx_v in enumerate(indices):
            if th_mask[idx]:
                th_tasks[idx_v].num = idx + 1
                variant.add_task(th_tasks[idx_v])

        for idx, idx_v in enumerate(pr_mask):
            if idx_v:
                pr_tasks[idx_v].num = idx + 1
                variant.add_task(pr_tasks[idx])

        # save variant document
        variant.save_test(out_path)

        # get table with keys of the variant
        key_table = variant.form_key()

        # add an empty table and replace it with created table
        key.add_paragraph(f'\nВариант {v_c + 1}')
        key.add_table(0, 0)
        key_root.replace(key.tables[-1]._element, key_table._element)

    for p in key.paragraphs:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = docx.shared.Pt(14)
            run.bold = True
    key.save(f'{out_path}/key.docx')
